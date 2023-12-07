# First, make sure you've installed these libraries:
# pip install streamlit
# pip install langchain
# pip install python-docx
# pip install openai

import os
import io
from docx import Document
from docx.shared import Pt
import streamlit as st
from apikey import apikey
import langchain
from langchain.llms import openai
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SequentialChain
from langchain.memory import ConversationBufferMemory
from datetime import datetime  

# Set up the environment variable for OpenAI API key
os.environ['OPENAI_API_KEY'] = apikey

# st.set_page_config(page_title="PMO Document Generator", page_icon="🧊", layout="wide")

# Initialize Streamlit app
st.title("PMO Document Generator")
#st.write("Some text here")
#selected_date = st.date_input("Select a date", datetime.today())

####

# User input
prompt = st.text_input('Enter prompt for documentation here')

# Function to save response to Word document
def save_to_word(doc_title, doc_content):
    # Create a new Document
    doc = Document()
    doc.add_heading(doc_title, 0)

    # Add the content to the Document
    p = doc.add_paragraph('')
    runner = p.add_run(doc_content)
    runner.font.size = Pt(12)  # Optionally set the font size

    # Save the document to an in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Return the buffer
    return buffer

# Prompt templates
#prompt tempalate
title_template = PromptTemplate(
    input_variables= ['topic'],
    template= """Please generate a comprehensive Project Initiation Document for the topic {topic}.
    \\This document should thoroughly delineate the scope of the project and provide sample text for each of the listed sections. 
    \\The structure should follow standard project initiation documentation format and include the following sections: Introduction: Briefly introduce the topic and its significance. 
    \\ Project Scope: Define the boundaries of the project, including what is within and outside the scope. 
    \\ Objectives: State the primary goals and objectives the project aims to achieve. 
    \\ Project Deliverables: List and describe the expected outcomes of the project.
    \\ Assumptions & Constraints: Identify any assumptions made during the planning phase and any constraints that may affect project execution.  
    \\ Project Milestones: Outline key milestones and their expected completion dates.    
    \\ Requirements: Detail the essential requirements necessary for the projects success.    
    \\ Project Management: Describe the management structure, roles, and responsibilities.    
    \\ Acceptance Criteria: Specify the criteria for accepting deliverables as complete and satisfactory.    
    \\ Project Resources: Identify the resources (human, technical, financial) required for the project.    
    \\ Communication Plan: Explain the communication strategy for stakeholders involved in the project.    
    \\ For each section, please provide an example text that is representative of content that would typically be included in a real-world project initiation document for the specified topic. 
    \\ The example text should be relevant and realistic, serving as a robust template for each category where applicable. 
    \\ The final output should be a detailed and structured document that can serve as a solid example for initiating a project related to {topic}."""
) 
script_template = PromptTemplate(
    input_variables= ['title'],
    template=  """ Please generate a comprehensive Project Initiation Document text on this TITLE :{topic} using Azure Framework include images"""
)

# memory
# using to store the memory of the conversation
memory = ConversationBufferMemory(input_key='topic', memory_key='chat_history')

# Memory and LLM setup
llm = openai.OpenAI(temperature= 0)
title_chain = LLMChain(llm=llm, prompt=title_template, verbose=True, output_key= 'title', memory=memory)
script_chain = LLMChain(llm=llm, prompt=script_template, verbose=True, output_key='script', memory=memory)
sequential_chain = SequentialChain(chains=[title_chain,script_chain], input_variables=['topic'], output_variables=['title', 'script'], verbose=True)

# Run chains if prompt is given
if prompt:
    # Get the response from the chains
    response = sequential_chain({'topic': prompt})
    
    # Display the response in the app
    st.write(response['title'])
    st.write(response['script'])

    # Allow user to download the response as a Word document
    combined_text = f"{response['title']}\n\n{response['script']}"
    word_file = save_to_word("Project Document", combined_text)
    st.download_button(label='📝 Download Word Document',
                       data=word_file,
                       file_name='project_document.docx',
                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Expandable section for message history
    with st.expander('Message History'):
        st.info(memory.buffer) 