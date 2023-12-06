# First, make sure you've installed these libraries:
# pip install streamlit
# pip install langchain
# pip install python-docx
# pip install openai

import os
import io
from docx import Document
from docx.shared import Pt
import streamlit as st
from apikey import apikey
from langchain.llms import openai
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SequentialChain
from langchain.memory import ConversationBufferMemory
from datetime import datetime  

# Set up the environment variable for OpenAI API key
os.environ['OPENAI_API_KEY'] = apikey

# Initialize Streamlit app
st.title("PMO Document Generator")

# User input
prompt = st.text_input('Enter prompt for documentation here')

# Function to save response to Word document
def save_to_word(doc_title, doc_content):
    # Create a new Document
    doc = Document()
    doc.add_heading(doc_title, 0)

    # Add the content to the Document
    p = doc.add_paragraph('')
    runner = p.add_run(doc_content)
    runner.font.size = Pt(12)  # Optionally set the font size

    # Save the document to an in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Return the buffer
    return buffer

# Prompt templates
sow_template = PromptTemplate(
    input_variables={'project_title'},
    template="""Please generate a comprehensive Statement of Work (SoW) for the project titled {project_title}.
    \\This document should thoroughly delineate the scope of the project and provide sample text for each of the listed sections. 
    \\The structure should follow standard SoW format and include the following sections: 
Project Title and Introduction: Introduce the project's title and significance.
    \\Project Title: {project_title}
    \\SOW Prepared by: (Input name here)
    \\Client Name: (Input Client's name here)
    \\Background: Briefly describe the project and why it is being undertaken.
    \\Project Objectives: State the main goals and intended outcomes of this project.
    \\Scope of Work & Risk Identification: Define the parameters and boundaries of the project including risk identification pertinent to these elements.
    \\Tasks: Outline the main tasks and responsibilities that will be performed during the project.
    \\Project Deliverables: Describe the expected outcomes and end-products of the project.
    \\Risk Review Periods: Schedule for regular risk reviews throughout the duration of the project.
    \\Project Timeline: State the start, end, and key milestones dates.
    \\Payment Schedule: Outline the payment terms, milestones linked to payments.
    \\Risk Management: Describe your approach to risk management and how potential risks will be forecasted, addressed, and mitigated during the project.
    \\Terms and Conditions: Point out any legal agreements related to the project.
    \\Acceptance: Define what defines successful deliverables and outcome of the project.
    \\Each section should contain relevant and realistic content serving as a template for a comprehensive Project Statement of Work (SoW) on {project_title}.
    """
)
script_template = PromptTemplate(
    input_variables=['title'],
    template=""" Please generate a comprehensive Project Initiation Document text on this TITLE :{topic} using Azure Framework include images"""
)

# Memory
# using to store the memory of the conversation
memory = ConversationBufferMemory(input_key='topic', memory_key='chat_history')

# Memory and LLM setup
llm = openai.OpenAI(temperature=0.5)  # Adjust temperature as needed
title_chain = LLMChain(llm=llm, prompt=sow_template, verbose=True, output_key='title', memory=memory)
script_chain = LLMChain(llm=llm, prompt=script_template, verbose=True, output_key='script', memory=memory)
sequential_chain = SequentialChain(chains=[title_chain, script_chain], input_variables=['topic'], output_variables=['title', 'script'], verbose=True)

# Run chains if prompt is given
if prompt:
    # Get the response from the chains
    response = sequential_chain({'topic': prompt})
    
    # Display the response in the app
    st.write(response['title'])
    st.write(response['script'])

    # Allow the user to download the response as a Word document
    combined_text = f"{response['title']}\n\n{response['script']}"
    word_file = save_to_word("Project Document", combined_text)
    st.download_button(label='📝 Download Word Document',
                       data=word_file,
                       file_name='project_document.docx',
                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Expandable section for message history
    with st.expander('Message History'):
        st.info(memory.buffer)
